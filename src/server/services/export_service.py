"""行程导出服务 — 支持 Markdown / PDF / DOCX / 纯文本

从 Trip 的 itinerary_json（或 raw_markdown）生成多格式导出文件。
PDF 导出需要 weasyprint（可选依赖，Windows 上可能需要额外配置）。
"""

from io import BytesIO

import markdown as md_lib

from ..models import Trip


class ExportService:
    """行程多格式导出服务"""

    def __init__(self, trip: Trip):
        self.trip = trip
        self.md_content = trip.raw_markdown or self._build_markdown()

    # ========== 内部：Markdown 构建 ==========

    def _build_markdown(self) -> str:
        """从 itinerary_json 重建格式化的 Markdown"""
        lines = [
            f"# 🗺️ {self.trip.title}",
            f"",
            f"**目的地**：{self.trip.city}  |  **天数**：{self.trip.days}天  |  **生成时间**：{self.trip.created_at.strftime('%Y-%m-%d %H:%M') if self.trip.created_at else ''}",
            f"",
            f"---",
            f"",
        ]

        itinerary = self.trip.itinerary_json or {}
        if isinstance(itinerary, dict):
            for day_label, spots in itinerary.items():
                lines.append(f"## {day_label}")
                lines.append("")
                if isinstance(spots, list):
                    for s in spots:
                        if isinstance(s, dict):
                            time = s.get("time", "")
                            name = s.get("name", "")
                            duration = s.get("duration", "")
                            price = s.get("price", 0)
                            transport = s.get("transport", "")
                            next_spot = s.get("next_spot", "")
                            next_distance = s.get("next_distance", "")

                            lines.append(f"### ⏰ {time}  —  {name}")
                            lines.append("")
                            if duration:
                                lines.append(f"- 🕐 建议游玩：**{duration}**")
                            if price:
                                lines.append(f"- 🎫 门票：**¥{price}**")
                            if transport:
                                transport_line = f"- 🚗 交通：{transport}"
                                if next_spot:
                                    transport_line += f" → {next_spot}"
                                if next_distance:
                                    transport_line += f"（{next_distance}）"
                                lines.append(transport_line)
                            lines.append("")
                lines.append("")

        if not itinerary:
            lines.append("> 暂无详细行程数据")
            lines.append("")

        lines.append("---")
        lines.append("")
        lines.append("> 📌 本行程由 **TripAgent AI** 自动生成，仅供参考。实际出行请以当地最新信息为准。")

        return "\n".join(lines)

    # ========== 各格式导出 ==========

    def to_markdown(self) -> bytes:
        """返回 Markdown 文件 bytes"""
        return self.md_content.encode("utf-8")

    def to_html(self) -> str:
        """Markdown → 带排版样式的 HTML"""
        md = md_lib.Markdown(extensions=["tables", "fenced_code", "codehilite"])
        body = md.convert(self.md_content)
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{self.trip.title} — TripAgent</title>
  <style>
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Microsoft YaHei", sans-serif;
      max-width: 800px;
      margin: 0 auto;
      padding: 2em 1.5em;
      color: #1a1a2e;
      line-height: 1.8;
      background: #faf8f5;
    }}
    h1 {{ color: #1a1a2e; font-size: 2em; border-bottom: 3px solid #e07a5f; padding-bottom: 0.3em; margin-bottom: 0.5em; }}
    h2 {{ color: #3d405b; font-size: 1.4em; margin: 1.8em 0 0.8em; padding: 0.5em 0; border-bottom: 1px solid #e0d5c7; }}
    h3 {{ color: #e07a5f; font-size: 1.1em; margin: 1.2em 0 0.4em; }}
    p {{ margin: 0.5em 0; }}
    ul {{ margin: 0.4em 0 0.4em 1.2em; }}
    li {{ margin: 0.3em 0; }}
    blockquote {{ border-left: 4px solid #81b29a; padding: 0.5em 1em; margin: 1.5em 0; color: #6b705c; background: #f4f1ea; border-radius: 0 8px 8px 0; }}
    strong {{ color: #e07a5f; }}
  </style>
</head>
<body>
{body}
</body>
</html>"""

    def to_pdf(self) -> bytes:
        """HTML → PDF

        优先使用 weasyprint，不可用时降级为返回 HTML（前端可用浏览器打印）。
        """
        try:
            from weasyprint import HTML
            return HTML(string=self.to_html()).write_pdf()
        except ImportError:
            # weasyprint 不可用，返回 HTML 内容并标记为 PDF fallback
            html = self.to_html()
            return html.encode("utf-8")

    def to_docx(self) -> bytes:
        """生成 Word 文档"""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        # 标题
        title_para = doc.add_heading(self.trip.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"目的地：{self.trip.city}  |  {self.trip.days}天").italic = True
        doc.add_paragraph("")  # 空行

        # 每日行程
        itinerary = self.trip.itinerary_json or {}
        if isinstance(itinerary, dict):
            for day_label, spots in itinerary.items():
                doc.add_heading(day_label, level=1)
                if isinstance(spots, list):
                    for s in spots:
                        if isinstance(s, dict):
                            p = doc.add_paragraph()
                            time = s.get("time", "")
                            name = s.get("name", "")
                            p.add_run(f"⏰ {time}  ").bold = False
                            p.add_run(name).bold = True
                            duration = s.get("duration", "")
                            price = s.get("price", 0)
                            if duration or price:
                                detail = doc.add_paragraph(style="List Bullet")
                                if duration:
                                    detail.add_run(f"建议游玩：{duration}  ")
                                if price:
                                    detail.add_run(f"门票：¥{price}")
                            transport = s.get("transport", "")
                            if transport:
                                tp = doc.add_paragraph(style="List Bullet")
                                tp.add_run(f"交通：{transport}")
                                if s.get("next_spot"):
                                    tp.add_run(f" → {s['next_spot']}")
                            doc.add_paragraph("")

        # 页脚声明
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.add_run("📌 本行程由 TripAgent AI 自动生成，仅供参考。").italic = True

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def to_plain_text(self) -> str:
        """生成无格式纯文本，供前端一键复制"""
        lines = [
            f"{self.trip.title}",
            f"目的地：{self.trip.city}  ·  {self.trip.days}天",
            f"{'─' * 50}",
        ]

        itinerary = self.trip.itinerary_json or {}
        if isinstance(itinerary, dict):
            for day_label, spots in itinerary.items():
                lines.append(f"\n【{day_label}】")
                if isinstance(spots, list):
                    for s in spots:
                        if isinstance(s, dict):
                            time = s.get("time", "")
                            name = s.get("name", "未知景点")
                            price_str = f"  🎫¥{s['price']}" if s.get("price") else ""
                            duration_str = f"  ({s['duration']})" if s.get("duration") else ""
                            lines.append(
                                f"  {time}  {name}{duration_str}{price_str}"
                            )
                            if s.get("transport"):
                                lines.append(f"         🚗 {s['transport']} → {s.get('next_spot', '下一站')}")

        lines.append(f"\n{'─' * 50}")
        lines.append("📌 由 TripAgent AI 自动生成 · 仅供参考")
        return "\n".join(lines)

    # ========== 文件工具 ==========

    @staticmethod
    def get_mime_type(format: str) -> str:
        """获取格式对应的 MIME 类型"""
        return {
            "md": "text/markdown; charset=utf-8",
            "html": "text/html; charset=utf-8",
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "txt": "text/plain; charset=utf-8",
        }.get(format, "application/octet-stream")

    @staticmethod
    def get_extension(format: str) -> str:
        """获取格式对应的文件扩展名"""
        return {"md": "md", "html": "html", "pdf": "pdf", "docx": "docx", "txt": "txt"}.get(
            format, "bin"
        )

    def get_filename(self, format: str) -> str:
        """生成下载文件名"""
        ext = self.get_extension(format)
        safe_title = self.trip.title.replace(" ", "_").replace("/", "-")[:30]
        safe_city = (self.trip.city or "trip").replace(" ", "_")
        return f"{safe_title}_{safe_city}.{ext}"
